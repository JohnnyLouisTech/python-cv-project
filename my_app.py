from pydoc import doc
from turtle import title
from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text): pyttsx3.speak(text)


document = Document()

# profile picture section
document.add_picture('johnny1.jpg', width=Inches(1.0))

# name, phone number, email, details, etc
name = input('What is your name? ')
speak('Hello ' + name + ' how are you doing today? ')
city = input('Where do you live? ')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')
website = input('What is your website? ')

document.add_paragraph(
    name + ' | ' + city + ' | ' + phone_number + ' | ' + email + ' | ' + website)

# about me section
document.add_heading('Abour me')
document.add_paragraph(
    input('Tell us  about yourself? ')
)
# work experience section
document.add_heading('Work Experience')
p = document.add_paragraph()

job_title = input('Enter job title: ')
company = input('Enter company name: ')
from_date = input('From date: ')
to_date = input('To Date: ')

p.add_run(job_title + '-' + company + '\n').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your work experience at ' + company + ' : ')
p.add_run(experience_details)

# work experience section continue
while True:
    has_more_experiences = input(
        'Do you have any more work experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        job_title = input('Enter job title: ')
        company = input('Enter company name: ')
        from_date = input('From date: ')
        to_date = input('To Date: ')

        p.add_run(job_title + '-' + company + '\n').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your work experience at ' + company + ':')
        p.add_run(experience_details)

    else:
        break

# education section
document.add_heading('Education')
p = document.add_paragraph()

school = input('Enter school name: ')
from_date = input('From date: ')
to_date = input('To Date: ')

p.add_run(school + '-').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

school_details = input(
    'Describe your education at ' + school + ':')
p.add_run(school_details)

# Education secion continue......
while True:
    has_more_educations = input(
        'Do you have any more educations? Yes or No ')
    if has_more_educations.lower() == 'yes':
        p = document.add_paragraph()

        school = input('Enter school name: ')
        from_date = input('From date: ')
        to_date = input('To Date: ')

        p.add_run(school + '-').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        school_details = input(
            'Describe your education at ' + school + ':')
        p.add_run(school_details)

    else:
        break


# list of skills
document.add_heading('Skills ')
skill = input('Enter skill: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do your have any more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill:')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# cv footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV created by JohnnyLouisTech as a Python Project"

document.save('cv.docx')
